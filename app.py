from flask import Flask, abort, render_template, request, send_file , session, redirect, url_for, flash

from email_validator import validate_email, EmailNotValidError
from flask_login import login_user, logout_user, login_required, current_user
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager
from flask_login import UserMixin
from werkzeug.security import generate_password_hash, check_password_hash
import pandas as pd
from datetime import datetime
import json
from docx import Document
from docx.shared import Inches , Cm , Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT , WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import date
from openpyxl import load_workbook



app = Flask(__name__, static_folder='static', static_url_path='/category/static')
# app.secret_key = 'your_secret_key_here'




#configure SQL Alchamy
app.config.update(
    SECRET_KEY = 'your‑very‑secret‑key',
    SQLALCHEMY_DATABASE_URI = 'sqlite:///users.db',  # or your DB of choice
    SQLALCHEMY_TRACK_MODIFICATIONS = False
)

db = SQLAlchemy(app)


#user Class 
class CustomUser(db.Model, UserMixin):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(128), nullable=False)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)


    def check_password(self, password):
        return check_password_hash(self.password_hash, password)
    





VALID_TEMPLATES = {
    'assets':          'assets.html',
    'create-asset':    'create-asset.html',
    'create-employee': 'create-employee.html',
}
@app.route('/')
def index():
     if 'username' in session:
        return render_template('index.html')
     else:
        return redirect(url_for('login'))





# Login
@app.route('/login', methods=['GET','POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        
        user = CustomUser.query.filter_by(username=username).first()
        if user and user.check_password(password):
           
            flash('Logged in successfully.', 'success')
            session['username'] = username
            return render_template('index.html')
        flash('Invalid email or password.', 'danger')
        return "Invalid User"

    return render_template('login.html')

@app.route('/debug-session')
def debug_session():
    return dict(session)  # or jsonify(session)

# Register
# @app.route('/register', methods=['GET','POST'])
# def register():
#     if request.method == 'POST':
#         email = request.form['email']
#         pw = request.form['password']
#         # validate email
#         try:
#             valid = validate_email(email)
#             email = valid.email
#         except EmailNotValidError as e:
#             flash(str(e), 'danger')
#             return redirect(url_for('register'))

#         if CustomUser.query.filter_by(email=email).first():
#             flash('Email already registered.', 'warning')
#             return redirect(url_for('register'))

#         user = CustomUser(email=email)
#         user.set_password(pw)
#         db.session.add(user)
#         db.session.commit()
#         flash('Registration successful—please log in.', 'success')
#         return redirect(url_for('login'))

#     return render_template('register.html')


# @app.route('/dashboard')
# @login_required
# def dashboard():
#     if "admin" in session:
#         return render_template("index.html")
#     else:
#        render_template("login.html")

    

# @app.route('/logout')
# @login_required
# def logout():
#     logout_user()
#     flash('You’ve been logged out.', 'info')
#     return redirect(url_for('login'))








@app.route('/submit_category', methods=['POST'])


@app.route('/category/<category>')
def show_category(category):
    template = VALID_TEMPLATES.get(category)
    if not template:
        abort(404)
    return render_template(template)



def lookup_inventory():
    df = pd.read_excel("Inventory.xlsx")
    df.columns = df.columns.str.strip()
    
    results = df[[
        "Device Type", "Description", "S/N",
        "Department", "Previous User Name", "Condition"
    ]]
    devices = results.to_dict(orient="records")

    return render_template("main-inventory.html", devices=devices)






def lookup_employees():
    df = pd.read_excel("Assetes & Custody.xlsx")
    df["Employee ID"].fillna(method='ffill', inplace=True)
    df.fillna("", inplace=True)

    # Include employee and device info
    devices = df[["Employee Name", "Employee ID", "Department", "Device Type", "Description", "S/N"]].to_dict(orient='records')
    
    return render_template("active-employees.html",devices=devices)
CATEGORY_HANDLERS = {
    'inventory': lookup_inventory,
    'employees': lookup_employees,
}
@app.route('/submit_category_lookup/category/<category>')
def submit_category_lookup(category):
   handler = CATEGORY_HANDLERS.get(category)
   if not handler:
        abort(404)
   return handler()


@app.route('/search')
def search_inventory():
    sn_query = request.args.get('sn', '').strip().lower()

    df = pd.read_excel("Inventory.xlsx")
    df.columns = df.columns.str.strip()
    df["S/N"] = df["S/N"].astype(str)

    if sn_query:
        df = df[df["S/N"].str.lower().str.contains(sn_query)]

    results = df[[
        "Device Type", "Description", "S/N",
        "Department", "Previous User Name", "Condition"
    ]]
    devices = results.to_dict(orient="records")
    return render_template("main-inventory.html", devices=devices, sn_query=sn_query)


@app.route('/search_employee')
def search_employee():
    emp_query = request.args.get('emp', '').strip().lower()

    df = pd.read_excel("Assetes & Custody.xlsx")
    df.columns = df.columns.str.strip()
    df["Employee ID"] = df["Employee ID"].astype(str)

    if emp_query:
        df = df[df["Employee ID"].str.lower().str.contains(emp_query)]

    results = df[[
        "Employee Name", "Employee ID" ,"Department","Device Type" ,"Description", "S/N"
        
    ]]
    devices = results.to_dict(orient="records")
    return render_template("active-employees.html", devices=devices, emp_query=emp_query)

# create new employee 
@app.route('/create_employee', methods=['POST'])
def create_employee():
    session['id'] = request.form['id']  # Get input from form
    session['name'] = request.form['name']  # Get input from form
    session['department'] = request.form['department']  # Get input from form
    df = pd.read_excel("Inventory.xlsx")
    devices = df[["Device Type", "Description", "S/N","Department", "Previous User Name", "Condition"]].to_dict(orient="records")
    return render_template("inventory.html", devices = devices)
    








# create new asset  
@app.route('/create_asset', methods=['POST'])
def create_asset():
    print(request.values.get('type'))
    device_type = request.form.get('type')
    device_des = request.form.get('description')
    device_serial = request.form.get('serial')
    device_department = request.form.get('department')

    new_row={
        "Department": device_department,
        "Device Type":  device_type,
        "Description": device_des,
        "S/N": device_serial,
        "Previous User Name": "N/A",
        "Location":"IT",
        "Condition": "New"


    }
  
    df = pd.read_excel("Inventory.xlsx")
    df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)

    # Save back to Excel
    df.to_excel("Inventory.xlsx",index=False)
    

    
    
    results = df[["Device Type", "Description", "S/N","Department", "Previous User Name", "Condition"]]
    devices = results.to_dict(orient="records")
    return render_template("main-inventory.html", devices = devices)
   


# open employee items

@app.route('/submit', methods=['POST'])
def submit():
    session['id'] = request.form['emp_id']  # Get input from form
    df = pd.read_excel("Assetes & Custody.xlsx")
    df["Employee ID"].fillna(method='ffill', inplace=True)
    if not df[df["Employee ID"] == int(session.get('id'))].empty:
        result = df[df["Employee ID"] == int(session.get('id'))]
        devices = result[["Device Type", "Description", "S/N"]].to_dict(orient='records')
        session['name'] = result["Employee Name"].to_list()[0]
        session['department']= result["Department"].to_list()[0]
        return render_template('form.html', name = session.get('name') , emp_department = session['department'] ,emp_id = session.get('id'), assets=devices)
    
    return render_template('create-employee.html')


    
    



# open inventory

@app.route('/inventory', methods=['POST'])
def form():
    df = pd.read_excel("Inventory.xlsx")
    df.columns = df.columns.str.strip()
    # df["Device Type"].fillna(method='ffill', inplace=True)
    action = request.form.get('action')
    if action == 'receiving':
     devices = df[["Device Type", "Description", "S/N","Department", "Previous User Name", "Condition"]].to_dict(orient="records")
     return render_template("inventory.html", devices = devices)
   # device_data = devices[["Device Type", "Description", "S/N","Previous User Name", "Condition"]].to_dict(orient='records')
    elif action == 'handover':
             df = pd.read_excel("Assetes & Custody.xlsx")
             result = df[df["Employee ID"] == int(session.get('id'))]
             devices = result[["Device Type", "Description", "S/N"]].to_dict(orient='records')
             return render_template("generated_form.html", name = session.get('name') , emp_department = session['department'], emp_id = session.get('id') ,  devices = devices)

        
    

    
   

        




@app.route('/submit_receiving_form', methods=['POST'])
def submit_receiving_form():
    df = pd.read_excel("Assetes & Custody.xlsx")
    df.columns = df.columns.str.strip()
    # Get the employee information
    emp_name = session.get('name')
    emp_id = session.get('id')
    emp_department = session['department']
    receive_date = date.today()
    
    # Get the selected devices (convert from JSON string back to dictionary)
    selected_raw = request.form.getlist('selected_devices')
    selected_devices = [json.loads(device) for device in selected_raw]


 


  

    # df = pd.read_excel("Assetes & Custody.xlsx",sheet_name="Assets")
    
    # result = df[df["Employee ID"] == int(session.get('id'))]




    

    # Create a Word document
    doc = Document()
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture('sources\mel-logo.png')
    
    heading = doc.add_heading(level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.add_run('RECEIVING FORM')
    run.bold = True
    run.underline = True
    font = run.font
    font.name = 'Bookman Old Style'          # Set font name
    font.size = Pt(26)
    run.font.bold = True           # Set font size (points)
    font.color.rgb = RGBColor(0, 0, 0)
        
    

    # Add employee information
    paragraph = doc.add_paragraph()
    paragraph.add_run('Name:   \t').font.size = Pt(14)
    
    paragraph.add_run(emp_name).font.size = Pt(14)
    paragraph = doc.add_paragraph()
    paragraph.add_run('ID:   \t\t').font.size = Pt(14)
    paragraph.add_run(emp_id).font.size = Pt(14)
    paragraph = doc.add_paragraph()
    paragraph.add_run(f'Dept: \t\t{emp_department}').font.size = Pt(14)
    paragraph = doc.add_paragraph()
    paragraph.add_run(f'Date: \t\t{receive_date}').font.size = Pt(14)
    
    doc.add_paragraph('\n')
    

    # Add a table for the devices
    table = doc.add_table(rows=1, cols=2)
   
    table.autofit = True
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells

    hdr_cells[0].width = Inches(7)
    paragraph = hdr_cells[0].paragraphs[0]
    run = paragraph.add_run('Description')
    run.bold = True
    font = run.font
    run.font.bold = True
    font.name = 'Bookman Old Style'          # Set font name
    font.size = Pt(14)           # Set font size (points)
    font.color.rgb = RGBColor(0, 0, 0)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    hdr_cells[1].width = Inches(1)
    paragraph =  hdr_cells[1].paragraphs[0]
    run = paragraph.add_run('Condition')
    run.bold = True
    font.name = 'Bookman Old Style'          # Set font name
    font.size = Pt(14)  
    run.font.bold = True         # Set font size (points)
    font.color.rgb = RGBColor(0, 0, 0)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
 
    

    # Add the devices to the table
    for index,  device in enumerate(selected_devices):
        row_cells = table.add_row().cells
        row_cells[0].text = f"{index+1}. {device['Device Type']} {device['Description']}, S/N: {device['S/N']}"
       
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = row_cells[0].paragraphs[0]
        # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        row_cells[1].text = device['Condition']
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = row_cells[1].paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    row_cells = table.add_row().cells  
    paragraph =  row_cells[0].paragraphs[0]
    run = paragraph.add_run("**** nothing Follows ****" )
    run.bold = True


    
    # add signature box 
    doc.add_paragraph('\n')
   
    
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture('static/images/footer + ack.png', width=Inches(6), height=Inches(3.2))

    # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    # doc.add_paragraph('\n')
    
    
    #add acknowledgements boxes

        # table = doc.add_table(rows=1, cols=4)
        
        # # Access the two cells
        # cell1 = table.rows[0].cells[0]
        # cell2 = table.rows[0].cells[2]
    
        # # Add pictures into the cells
        # cell1.paragraphs[0].add_run().add_picture('static\\images\\acknowledgement-box-en.png', width=Inches(2.5), height=Inches(1.4))
        # cell2.paragraphs[0].add_run().add_picture('static\\images\\acknowledgement-box-ar.png', width=Inches(2.5), height=Inches(1.4))
        # cell1.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        # cell2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

   
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture('static\\images\\footer.png', width=Inches(7))

    # Save the document to a BytesIO object
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)





    # assign devices to employee in the master file
    
    new_row = {}

    if len(selected_devices) == 1:
        new_row['Employee Name'] =  emp_name
        new_row['Employee ID'] =  int(emp_id)
        new_row['Department'] = emp_department 
        new_row['Device Type'] =   selected_devices[index]['Device Type']
        new_row['Description'] =  selected_devices[index]["Description"]
        new_row['S/N'] =  selected_devices[index]['S/N']

    elif  len(selected_devices) > 1:   
        for index ,device in enumerate(selected_devices):
            new_row[index]['Employee Name'] =  emp_name
            new_row[index]['Employee ID'] = int(emp_id)
            new_row[index]['Department'] = emp_department 
            new_row[index]['Device Type'] =   selected_devices[index]['Device Type']
            new_row[index]['Description'] =  selected_devices[index]["Description"]
            new_row[index]['S/N'] =  selected_devices[index]['S/N']

   
    
  
   
    df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)

    # Save back to Excel
    df.to_excel("Assetes & Custody.xlsx",index=False)

    # Load the workbook and sheet
    wb = load_workbook("Assetes & Custody.xlsx")
    ws = wb.active  

    # Auto-adjust column widths
    for column_cells in ws.columns:
        max_length = 0
        column = column_cells[0].column_letter  
        for cell in column_cells:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[column].width = adjusted_width

    # Save again
    wb.save("Assetes & Custody.xlsx")


    # delete devices from inventory

    df = pd.read_excel("Inventory.xlsx")
    removed_device_serial = selected_devices[0]["S/N"]
    df = df[df['S/N'] != removed_device_serial]

    df.to_excel('Inventory.xlsx', index=False)
    wb = load_workbook("Inventory.xlsx")
    ws = wb.active  

    # Auto-adjust column widths
    for column_cells in ws.columns:
        max_length = 0
        column = column_cells[0].column_letter  
        for cell in column_cells:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[column].width = adjusted_width

    # Save again
    wb.save("Inventory.xlsx")

    

    


    # Send the file as a response
    return send_file(doc_stream, as_attachment=True, download_name=F"receiving_form_{emp_id}_{emp_name[0:2]}.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


    


@app.route('/submit_handover_form', methods=['POST'])
def submit_handover_form():
  
    # Get the selected devices (convert from JSON string back to dictionary)
    selected_raw = request.form.getlist('selected_devices')
    if  selected_raw:
        selected_devices = [json.loads(device) for device in selected_raw]
        
        # Get the employee information
        emp_name = session.get('name')
        emp_id = session.get('id')
        emp_department = session['department']
        handover_date = date.today()


    


        # add devices to inventory


        df = pd.read_excel("Inventory.xlsx")
        new_row={
            "Previous User Name": emp_name,
            "Previous User EN#": int(emp_id),
            "Device Type": selected_devices[0]['Device Type'],
            "Description": selected_devices[0]["Description"],
            "S/N": selected_devices[0]['S/N']
        }
    
    
        df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)

        # Save back to Excel
        df.to_excel("Inventory.xlsx",index=False)

        # Load the workbook and sheet
        wb = load_workbook("Inventory.xlsx")
        ws = wb.active  

        # Auto-adjust column widths
        for column_cells in ws.columns:
            max_length = 0
            column = column_cells[0].column_letter  
            for cell in column_cells:
                try:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                except:
                    pass
            adjusted_width = (max_length + 2)
            ws.column_dimensions[column].width = adjusted_width

        # Save again
        wb.save("Inventory.xlsx")



        # delete devices of employee from the master file

        df = pd.read_excel("Assetes & Custody.xlsx")
        for device in selected_devices :
         removed_device_serial = device["S/N"]

        df = df[df['S/N'] != removed_device_serial]

        df.to_excel('Assetes & Custody.xlsx', index=False)
        # Load the workbook and sheet
        wb = load_workbook("Assetes & Custody.xlsx")
        ws = wb.active  

        # Auto-adjust column widths
        for column_cells in ws.columns:
            max_length = 0
            column = column_cells[0].column_letter  
            for cell in column_cells:
                try:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                except:
                    pass
            adjusted_width = (max_length + 2)
            ws.column_dimensions[column].width = adjusted_width

        # Save again
        wb.save("Assetes & Custody.xlsx")



        # Create a Word document
        doc = Document()
        section = doc.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture('sources\mel-logo.png')
        
        heading = doc.add_heading(level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = heading.add_run('HANDOVER FORM')
        run.bold = True
        run.underline = True
        font = run.font
        font.name = 'Bookman Old Style'          # Set font name
        font.size = Pt(26)  
        run.font.bold = True            # Set font size (points)
        font.color.rgb = RGBColor(0, 0, 0)
        

        # Add employee information
        paragraph = doc.add_paragraph()
        paragraph.add_run('Name:   \t').font.size = Pt(14)
    
        paragraph.add_run(emp_name).font.size = Pt(14)
        paragraph = doc.add_paragraph()
        paragraph.add_run('ID:   \t\t').font.size = Pt(14)
        paragraph.add_run(emp_id).font.size = Pt(14)
        paragraph = doc.add_paragraph()
        paragraph.add_run(f'Dept: \t\t{emp_department}').font.size = Pt(14)
        paragraph = doc.add_paragraph()
        paragraph.add_run(f'Date: \t\t{handover_date}').font.size = Pt(14)
    
        doc.add_paragraph('\n')

        # Add a table for the devices
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
       
        hdr_cells[0].width = Inches(6)
        paragraph = hdr_cells[0].paragraphs[0]
        run = paragraph.add_run("List of Items")
        run.bold = True
        font.name = 'Bookman Old Style'          # Set font name
        font.size = Pt(14)           # Set font size (points)
        font.color.rgb = RGBColor(0, 0, 0)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        hdr_cells[1].width = Inches(1)
        paragraph = hdr_cells[1].paragraphs[0]
        run = paragraph.add_run("S/N")
        run.bold = True
        font.name = 'Bookman Old Style'          # Set font name
        font.size = Pt(14)  
        run.font.bold = True         # Set font size (points)
        font.color.rgb = RGBColor(0, 0, 0)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
       
        hdr_cells[2].width = Inches(1)
        paragraph = hdr_cells[2].paragraphs[0]
        run = paragraph.add_run("Remarks/Password")
        run.bold = True
        font.name = 'Bookman Old Style'          # Set font name
        font.size = Pt(14)  
        run.font.bold = True         # Set font size (points)
        font.color.rgb = RGBColor(0, 0, 0)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        
    
        
        # Add the devices to the table
        for index,  device in enumerate(selected_devices):
            row_cells = table.add_row().cells
            row_cells[0].text = f"{index+1}. {device['Description']}"
            # row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = row_cells[0].paragraphs[0]
            # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    
            row_cells[1].text = device['S/N']
            row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = row_cells[1].paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # row_cells[3].text = device['S/N']
            # row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # paragraph = row_cells[3].paragraphs[0]
            # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        
        row_cells = table.add_row().cells  
        row_cells[0].text = "Others:" 
           



        # add signature box 
        doc.add_paragraph('\n')
        doc.add_paragraph('\n')
        doc.add_paragraph('\n')
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture('static/images/handover-sign.png',width=Inches(6.2))


        footer = doc.sections[0].footer
        paragraph = footer.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture('static\\images\\footer.png', width=Inches(7))


        # Save the document to a BytesIO object
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        
        # Send the file as a response
        return send_file(doc_stream, as_attachment=True, download_name=F"handover_form_{emp_id}_{emp_name[0:2]}.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    
    return("No selected Devices")



# create users 

def seed_users():
     users = [
        {'username': 'admin', 'password': 'Pass123'},
        {'username': 'Mustafa', 'password': 'Mel@2025$'},
        {'username': 'Khalid', 'password': 'Mel@2024$'},
    ]

     for u in users:
        print(f"Checking if user exists: {u['username']}")
        if not CustomUser.query.filter_by(username=u['username']).first():
            print(f"Adding user: {u['username']}")
            user = CustomUser(username=u['username'])
            user.set_password(u['password'])
            db.session.add(user)
        else:
            print(f"User already exists: {u['username']}")

     try:
        db.session.commit()
        print("All users committed.")
     except Exception as e:
        print(f"Commit failed: {e}")
        db.session.rollback()





if __name__ == '__main__':
    with app.app_context():
        db.drop_all()
        db.create_all()
        seed_users()
    app.run(debug=True)